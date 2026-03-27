"""
upload_file.py — Canvas 3-step file upload process.

Step 1: Notify Canvas API → get pre-authorized S3 upload URL + params
Step 2: POST file binary directly to S3 (NO Authorization header)
Step 3: Confirm with Canvas API → get file_id

Returns file_id for use in assignment submission.

Usage (as module):
  from tools.upload_file import upload_file
  file_id = upload_file(client, course_id, assignment_id, file_path)
"""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from tools.canvas_client import CanvasClient

CONTENT_TYPE_MAP = {
    ".pdf":  "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".txt":  "text/plain",
    ".md":   "text/markdown",
    ".html": "text/html",
    ".py":   "text/x-python",
    ".csv":  "text/csv",
}


def upload_file(client, course_id, assignment_id, file_path):
    """
    Upload a file to Canvas for an assignment submission.
    Returns the Canvas file_id (int).
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    file_name = os.path.basename(file_path)
    file_size = os.path.getsize(file_path)
    ext = os.path.splitext(file_name)[1].lower()
    content_type = CONTENT_TYPE_MAP.get(ext, "application/octet-stream")

    print(f"[upload_file] Step 1: Notifying Canvas API...")
    # Step 1 — Notify Canvas, get upload URL
    notify_response = client.post(
        f"/courses/{course_id}/assignments/{assignment_id}/submissions/self/files",
        json={
            "name": file_name,
            "size": file_size,
            "content_type": content_type,
        },
    )

    upload_url = notify_response.get("upload_url")
    upload_params = notify_response.get("upload_params", {})

    if not upload_url:
        raise RuntimeError(f"Canvas did not return an upload URL: {notify_response}")

    print(f"[upload_file] Step 2: Uploading to S3...")
    # Step 2 — Upload binary directly to S3 (NO auth header — must use raw post)
    with open(file_path, "rb") as fh:
        file_data = fh.read()

    # S3 requires form fields from upload_params BEFORE the file field
    form_data = {k: (None, v) for k, v in upload_params.items()}
    form_data["file"] = (file_name, file_data, content_type)

    import requests
    resp = requests.post(upload_url, files=form_data, timeout=120)

    if resp.status_code not in (200, 201, 301, 302, 303):
        raise RuntimeError(f"S3 upload failed {resp.status_code}: {resp.text[:300]}")

    # S3 may redirect to a confirmation URL
    confirm_url = resp.headers.get("Location") or resp.url

    print(f"[upload_file] Step 3: Confirming with Canvas...")
    # Step 3 — Confirm with Canvas to get file_id
    confirm_resp = client._request("GET", confirm_url)
    file_data_resp = confirm_resp.json()
    file_id = file_data_resp.get("id")

    if not file_id:
        raise RuntimeError(f"Canvas confirmation did not return file ID: {file_data_resp}")

    print(f"[upload_file] Upload complete. File ID: {file_id}")
    return file_id


def save_response_as_file(response_text, assignment_id, run_id):
    """Save LLM response text as a formatted Word document (.docx)."""
    return save_response_as_docx(response_text, assignment_id, run_id)


def save_response_as_docx(response_text, assignment_id, run_id, title=None):
    """
    Convert LLM response (markdown-ish text) into a properly formatted .docx file.
    Styling: Calibri 11pt body, student-appropriate margins, headings for ## lines.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re

        doc = Document()

        # ── Page margins (1 inch all around — standard student paper) ──────────
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1)
            section.right_margin  = Inches(1)

        # ── Default style: Calibri 11pt ────────────────────────────────────────
        style = doc.styles["Normal"]
        font  = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # ── Title ──────────────────────────────────────────────────────────────
        if title:
            t = doc.add_heading(title, level=1)
            t.runs[0].font.color.rgb = RGBColor(0x1F, 0x39, 0x64)  # dark navy

        # ── Parse and render lines ─────────────────────────────────────────────
        lines = response_text.strip().splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]

            # Heading 2 (##)
            if line.startswith("## "):
                h = doc.add_heading(line[3:].strip(), level=2)
                h.runs[0].font.size = Pt(13)
                i += 1
                continue

            # Heading 3 (###)
            if line.startswith("### "):
                h = doc.add_heading(line[4:].strip(), level=3)
                h.runs[0].font.size = Pt(12)
                i += 1
                continue

            # Heading 1 (#)
            if re.match(r"^# [^#]", line):
                h = doc.add_heading(line[2:].strip(), level=1)
                i += 1
                continue

            # Horizontal rule (--- or ***)
            if re.match(r"^[-\*]{3,}$", line.strip()):
                doc.add_paragraph("─" * 60)
                i += 1
                continue

            # Bullet list item
            if re.match(r"^[\-\*\•] ", line):
                text = line[2:].strip()
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_formatting(p, text)
                i += 1
                continue

            # Numbered list item
            if re.match(r"^\d+\. ", line):
                text = re.sub(r"^\d+\. ", "", line).strip()
                p = doc.add_paragraph(style="List Number")
                _add_inline_formatting(p, text)
                i += 1
                continue

            # Empty line → paragraph break
            if line.strip() == "":
                i += 1
                continue

            # Normal paragraph — collect consecutive non-blank lines
            para_lines = []
            while i < len(lines) and lines[i].strip() != "" and not lines[i].startswith("#") and not re.match(r"^[\-\*\•\d]", lines[i]):
                para_lines.append(lines[i])
                i += 1
            if para_lines:
                p = doc.add_paragraph()
                _add_inline_formatting(p, " ".join(para_lines))

        out_dir   = f".tmp/responses_{run_id}"
        os.makedirs(out_dir, exist_ok=True)
        file_path = f"{out_dir}/{assignment_id}_submission.docx"
        doc.save(file_path)
        print(f"[upload_file] Saved Word document → {file_path}")
        return file_path

    except ImportError:
        # python-docx not installed — fall back to .txt
        print("[upload_file] Warning: python-docx not installed, falling back to .txt")
        out_dir   = f".tmp/responses_{run_id}"
        os.makedirs(out_dir, exist_ok=True)
        file_path = f"{out_dir}/{assignment_id}_submission.txt"
        with open(file_path, "w") as f:
            f.write(response_text)
        return file_path


def _add_inline_formatting(paragraph, text):
    """Add a run to paragraph with **bold** and *italic* markdown parsed."""
    import re
    # Split on **bold** and *italic* markers
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run       = paragraph.add_run(part[2:-2])
            run.bold  = True
        elif part.startswith("*") and part.endswith("*"):
            run        = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
